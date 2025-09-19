#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word转GitBook工具使用示例
"""

import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches

# 添加当前目录到Python路径
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from word_to_gitbook import WordToGitBookConverter, GitBookConfig


def create_sample_word_document():
    """创建一个示例Word文档用于测试"""
    print("创建示例Word文档...")
    
    doc = Document()
    
    # 添加标题
    doc.add_heading('示例技术文档', 0)  # 文档标题
    
    # 添加介绍段落
    intro = doc.add_paragraph('这是一个示例技术文档，演示Word转GitBook工具的功能。')
    intro.add_run('本文档包含多种元素：标题、段落、').bold = True
    intro.add_run('粗体文本、')
    intro.add_run('斜体文本、').italic = True
    intro.add_run('以及表格和图片。')
    
    # 第一章
    doc.add_heading('第一章：项目介绍', 1)
    doc.add_paragraph('本章介绍项目的基本信息和背景。')
    
    # 1.1 节
    doc.add_heading('1.1 项目背景', 2)
    doc.add_paragraph('随着技术的发展，文档管理变得越来越重要。Word作为最常用的文档编辑工具，与GitBook这样的现代文档平台之间需要一个转换桥梁。')
    
    # 1.1.1 小节
    doc.add_heading('1.1.1 技术选型', 3)
    doc.add_paragraph('我们选择Python作为开发语言，主要原因包括：')
    doc.add_paragraph('• 丰富的文档处理库\n• 良好的跨平台支持\n• 简洁的语法结构', style='List Bullet')
    
    # 1.1.2 小节
    doc.add_heading('1.1.2 设计目标', 3)
    doc.add_paragraph('设计目标是创建一个功能完整、易于使用的转换工具。')
    
    # 1.2 节
    doc.add_heading('1.2 功能特性', 2)
    
    # 添加表格
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '功能'
    hdr_cells[1].text = '描述'
    hdr_cells[2].text = '状态'
    
    # 添加表格数据
    features = [
        ('文本转换', '支持所有文本格式的转换', '✅ 完成'),
        ('图片提取', '自动提取并保存图片文件', '✅ 完成'),
        ('表格转换', '将Word表格转为Markdown格式', '✅ 完成'),
        ('目录生成', '根据标题自动生成目录结构', '✅ 完成'),
        ('配置文件', '生成GitBook所需的配置文件', '✅ 完成')
    ]
    
    for feature, desc, status in features:
        row_cells = table.add_row().cells
        row_cells[0].text = feature
        row_cells[1].text = desc
        row_cells[2].text = status
    
    # 第二章
    doc.add_heading('第二章：安装和使用', 1)
    doc.add_paragraph('本章详细说明工具的安装和使用方法。')
    
    # 2.1 节
    doc.add_heading('2.1 环境要求', 2)
    doc.add_paragraph('使用本工具需要满足以下环境要求：')
    doc.add_paragraph('• Python 3.7或更高版本\n• 必要的Python依赖包\n• 足够的磁盘空间存储输出文件')
    
    # 2.2 节
    doc.add_heading('2.2 安装步骤', 2)
    doc.add_paragraph('请按照以下步骤安装：')
    
    # 2.2.1 小节
    doc.add_heading('2.2.1 下载代码', 3)
    doc.add_paragraph('从GitHub仓库下载最新版本的代码。')
    
    # 2.2.2 小节
    doc.add_heading('2.2.2 安装依赖', 3)
    doc.add_paragraph('运行以下命令安装依赖：')
    
    # 添加代码样式段落
    code_para = doc.add_paragraph('pip install -r requirements.txt')
    code_para.style = 'Intense Quote'
    
    # 第三章
    doc.add_heading('第三章：高级功能', 1)
    doc.add_paragraph('本章介绍工具的高级功能和配置选项。')
    
    # 3.1 节
    doc.add_heading('3.1 目录配置', 2)
    doc.add_paragraph('工具支持灵活的目录级别配置，可以通过参数控制生成的目录深度。')
    
    # 3.2 节
    doc.add_heading('3.2 输出定制', 2)
    doc.add_paragraph('用户可以自定义输出目录结构和文件命名规则。')
    
    # 保存文档
    sample_file = Path('sample_document.docx')
    doc.save(sample_file)
    print(f"示例文档已创建: {sample_file}")
    return sample_file


def example_basic_conversion():
    """基本转换示例"""
    print("\n=== 基本转换示例 ===")
    
    # 创建示例文档
    sample_file = create_sample_word_document()
    
    # 基本配置
    config = GitBookConfig(
        title="示例技术文档",
        description="这是一个Word转GitBook的示例文档",
        max_toc_level=3,
        output_dir="example_basic_output"
    )
    
    # 执行转换
    converter = WordToGitBookConverter(config)
    converter.convert(str(sample_file))
    
    print("基本转换完成！")
    print(f"输出目录: {config.output_dir}")


def example_advanced_conversion():
    """高级转换示例"""
    print("\n=== 高级转换示例 ===")
    
    # 创建示例文档
    sample_file = create_sample_word_document()
    
    # 高级配置
    config = GitBookConfig(
        title="高级技术手册",
        description="包含完整功能演示的技术手册",
        language="zh-hans",
        max_toc_level=4,  # 显示4级目录
        output_dir="example_advanced_output",
        assets_dir="images"  # 自定义图片目录
    )
    
    # 执行转换
    converter = WordToGitBookConverter(config)
    converter.convert(str(sample_file))
    
    print("高级转换完成！")
    print(f"输出目录: {config.output_dir}")


def example_minimal_conversion():
    """最小配置转换示例"""
    print("\n=== 最小配置转换示例 ===")
    
    # 创建示例文档
    sample_file = create_sample_word_document()
    
    # 最小配置（只显示1级目录）
    config = GitBookConfig(
        title="简化文档",
        description="只显示主要章节的简化文档",
        max_toc_level=1,  # 只显示1级目录
        output_dir="example_minimal_output"
    )
    
    # 执行转换
    converter = WordToGitBookConverter(config)
    converter.convert(str(sample_file))
    
    print("最小配置转换完成！")
    print(f"输出目录: {config.output_dir}")


def cleanup_example_files():
    """清理示例文件"""
    print("\n=== 清理示例文件 ===")
    
    import shutil
    
    # 删除示例Word文档
    sample_file = Path('sample_document.docx')
    if sample_file.exists():
        sample_file.unlink()
        print(f"已删除: {sample_file}")
    
    # 删除输出目录
    output_dirs = [
        "example_basic_output",
        "example_advanced_output", 
        "example_minimal_output"
    ]
    
    for dir_name in output_dirs:
        dir_path = Path(dir_name)
        if dir_path.exists():
            shutil.rmtree(dir_path)
            print(f"已删除目录: {dir_path}")


def main():
    """主函数"""
    print("Word转GitBook工具使用示例")
    print("=" * 50)
    
    try:
        # 运行各种示例
        example_basic_conversion()
        example_advanced_conversion()
        example_minimal_conversion()
        
        print("\n" + "=" * 50)
        print("所有示例转换完成！")
        print("\n请查看以下输出目录：")
        print("- example_basic_output/     (基本转换，3级目录)")
        print("- example_advanced_output/  (高级转换，4级目录)")
        print("- example_minimal_output/   (最小转换，1级目录)")
        
        print("\n你可以进入任意目录查看转换结果：")
        print("cd example_basic_output")
        print("ls -la")
        
        # 询问是否清理文件
        response = input("\n是否清理示例文件？(y/N): ").strip().lower()
        if response in ['y', 'yes']:
            cleanup_example_files()
        
    except Exception as e:
        print(f"运行示例时出错: {e}")
        return 1
    
    return 0


if __name__ == "__main__":
    exit(main())
